from docx import Document
import sys

def parse_docx_to_pages(file_path, max_chars_per_page=1000):
    print(f"Processing file: {file_path}")
    try:
        doc = Document(file_path)
        print(f"Document loaded, paragraph count: {len(doc.paragraphs)}")
        pages = []
        current_page = []
        char_count = 0

        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                para_html = f'<p>{para.text}</p>'
                para_chars = len(para.text)
                if char_count + para_chars > max_chars_per_page and current_page:
                    pages.append(''.join(current_page))
                    current_page = [para_html]
                    char_count = para_chars
                else:
                    current_page.append(para_html)
                    char_count += para_chars
        if current_page:
            pages.append(''.join(current_page))
        print(f"Pages generated: {len(pages)}")
        return pages
    except Exception as e:
        print(f"Error in parse_docx_to_pages: {str(e)}")
        raise

if __name__ == '__main__':
    file_path = sys.argv[1]
    pages = parse_docx_to_pages(file_path)
    for page in pages:
        print(page)